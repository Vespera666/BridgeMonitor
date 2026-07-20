#!/usr/bin/env python3
"""生成对流换热系数公式文档 (DOCX + PDF)"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 字体工具：确保中文正常显示 ──
CN_FONT = '微软雅黑'
EN_FONT = 'Consolas'

def set_run_font(run, name=CN_FONT, size=None, bold=False, color=None):
    """安全设置 run 的字体（中文+西文回退）"""
    run.font.name = name
    run.font.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color

def add_cn_paragraph(doc, text, size=11, bold=False, alignment=None):
    """添加中文段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=Pt(size), bold=bold)
    return p

def add_cn_heading(doc, text, level=2):
    """添加中文标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=run.font.size, bold=True)
    return h
from fpdf import FPDF
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DPI = 200  # 高清渲染

# ============================================================
# 1. 用 matplotlib 渲染 LaTeX 公式为 PNG
# ============================================================
FORMULAS = [
    # (编号, 名称, LaTeX, 变量说明)
    ("1", "牛顿冷却定律",
     r"$q = h \cdot A \cdot (T_s - T_\infty)$",
     "q — 热流密度 [W]    h — 对流换热系数 [W/(m²·K)]    A — 换热面积 [m²]    T_s — 壁面温度 [K]    T_∞ — 流体温度 [K]"),

    ("2", "努塞尔数 (Nusselt Number)",
     r"$Nu = \frac{h \cdot L}{k}$",
     "h — 对流换热系数 [W/(m²·K)]    L — 特征长度 [m]    k — 流体导热系数 [W/(m·K)]"),

    ("3", "雷诺数 (Reynolds Number)",
     r"$Re = \frac{\rho \cdot v \cdot L}{\mu}$",
     "ρ — 流体密度 [kg/m³]    v — 流速 [m/s]    L — 特征长度 [m]    μ — 动力粘度 [Pa·s]"),

    ("4", "普朗特数 (Prandtl Number)",
     r"$Pr = \frac{c_p \cdot \mu}{k}$",
     "c_p — 定压比热容 [J/(kg·K)]    μ — 动力粘度 [Pa·s]    k — 导热系数 [W/(m·K)]"),

    ("5", "迪图斯-贝尔特关联式 (管内湍流)",
     r"$Nu = 0.023 \cdot Re^{0.8} \cdot Pr^{n}$",
     "加热流体时 n = 0.4，冷却流体时 n = 0.3    适用范围：Re > 10⁴, 0.7 < Pr < 160, L/D > 60"),

    ("6", "自然对流关联式",
     r"$Nu = C \cdot (Gr \cdot Pr)^{n}$",
     "C, n — 实验常数（取决于流态和几何形状）    Gr — 格拉晓夫数"),

    ("7", "格拉晓夫数 (Grashof Number)",
     r"$Gr = \frac{g \cdot \beta \cdot (T_s - T_\infty) \cdot L^{3}}{\nu^{2}}$",
     "g — 重力加速度 [m/s²]    β — 体积膨胀系数 [1/K]    ν — 运动粘度 [m²/s]"),
]


def render_formula(latex, filename):
    """用 matplotlib 内置 mathtext 渲染公式"""
    fig, ax = plt.subplots(figsize=(5.5, 0.8))
    ax.text(0.5, 0.5, latex, fontsize=18, ha='center', va='center',
            transform=ax.transAxes)
    ax.axis('off')
    fig.savefig(filename, dpi=DPI, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)


# 渲染所有公式
img_dir = os.path.join(OUTPUT_DIR, "_formula_images")
os.makedirs(img_dir, exist_ok=True)

for num, name, latex, _ in FORMULAS:
    fname = os.path.join(img_dir, f"formula_{num}.png")
    render_formula(latex, fname)
    print(f"  [OK] Rendered {num}. {name}")

# ============================================================
# 2. 生成 DOCX
# ============================================================
doc = Document()

# 封面标题
title = doc.add_heading('对流换热系数 — 核心公式', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_run_font(run, size=Pt(22), bold=True)

add_cn_paragraph(doc,
    '对流换热系数（Convective Heat Transfer Coefficient）是表征流体与固体壁面之间'
    '换热能力的关键参数，单位为 W/(m²·K)。以下为工程计算中最常用的无量纲准则数与关联式。',
    size=11)
doc.add_paragraph()

for num, name, latex, note in FORMULAS:
    # 公式名称
    add_cn_heading(doc, f'{num}. {name}', level=2)

    # 公式图片
    img_path = os.path.join(img_dir, f"formula_{num}.png")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(4.5))

    # LaTeX 源码（小字灰色）
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_src = p_src.add_run(f'LaTeX: {latex}')
    set_run_font(run_src, name=EN_FONT, size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

    # 变量说明
    p_note = doc.add_paragraph(note)
    for run in p_note.runs:
        set_run_font(run, size=Pt(9))

docx_path = os.path.join(OUTPUT_DIR, "对流换热公式.docx")
doc.save(docx_path)
print(f"\n[DONE] DOCX saved: {docx_path}")

# ============================================================
# 3. 生成 PDF
# ============================================================
class FormulaPDF(FPDF):
    def __init__(self):
        super().__init__('P', 'mm', 'A4')
        # 注册中文字体（使用 Windows 自带微软雅黑）
        font_path = "C:/Windows/Fonts/msyh.ttc"
        if os.path.exists(font_path):
            self.add_font('CJK', '', font_path, uni=True)
            self.add_font('CJK', 'B', font_path, uni=True)  # 复用为粗体
            self.font_family = 'CJK'
        else:
            self.font_family = 'Helvetica'

    def header(self):
        if self.page_no() == 1:
            self.set_font(self.font_family, 'B', 16)
            self.cell(0, 12, '对流换热系数 — 核心公式', align='C', new_x="LMARGIN", new_y="NEXT")
            self.ln(6)

    def footer(self):
        self.set_y(-15)
        self.set_font(self.font_family, '', 8)
        self.cell(0, 10, f'第 {self.page_no()} 页', align='C')

pdf = FormulaPDF()
pdf.set_auto_page_break(True, 20)
pdf.add_page()

# 简介
pdf.set_font(pdf.font_family, '', 10)
pdf.multi_cell(0, 6,
    '对流换热系数（Convective Heat Transfer Coefficient）是表征流体与固体壁面之间'
    '换热能力的关键参数，单位为 W/(m²·K)。以下为工程计算中最常用的无量纲准则数与关联式。')
pdf.ln(4)

for num, name, latex, note in FORMULAS:
    # 检查是否需要新页
    if pdf.get_y() > 220:
        pdf.add_page()

    # 标题
    pdf.set_font(pdf.font_family, 'B', 12)
    pdf.cell(0, 8, f'{num}. {name}', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # 公式图片
    img_path = os.path.join(img_dir, f"formula_{num}.png")
    if os.path.exists(img_path):
        img_w = 120  # mm
        x = (pdf.w - img_w) / 2
        pdf.image(img_path, x=x, w=img_w)
        pdf.ln(4)

    # LaTeX 源码
    pdf.set_font(pdf.font_family, '', 7)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 4, f'LaTeX: {latex}', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    # 变量说明
    pdf.set_font(pdf.font_family, '', 8)
    pdf.multi_cell(0, 4.5, note)
    pdf.ln(3)

pdf_path = os.path.join(OUTPUT_DIR, "对流换热公式.pdf")
pdf.output(pdf_path)
print(f"[DONE] PDF saved: {pdf_path}")

print("\n完成！")
